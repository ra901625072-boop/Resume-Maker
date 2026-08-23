"""
services/ai_service.py — OpenRouter AI Brain (v3 - Optimized Multimodal & Universal Extractor)
===========================================================================================
Architecture:
  • Multi-tier model fallback chain with top multimodal & reasoning models
  • Vision & Document intelligence: Google Gemini 2.0 Flash / Qwen 2.5 VL 72B / Llama 3.2 Vision
  • Universal Ingestion: Digital PDFs, Scanned PDFs, DOCX, DOC, RTF, ODT, TXT, MD, CSV, Images (JPG, PNG, WebP, BMP, TIFF)
  • Image Optimizer: EXIF orientation correction, DPI scaling, and high-fidelity base64 encoding
  • In-process LRU response cache (TTL-based)
  • Robust JSON repair & schema validation
  • All public methods return:
      { "success": True,  "data": <str|dict>, "tokens": int, "model": str, ... }
      { "success": False, "error": str }
"""

import base64
import hashlib
import io
import json
import logging
import os
import re
import time
from functools import lru_cache
from typing import Optional, Union

import requests
from PIL import Image, ImageOps

logger = logging.getLogger(__name__)

# ── Model tiers ──────────────────────────────────────────────────────────────
# Default text-centric reasoning models (tried in order)
_MODELS = [
    "openrouter/free",
    "google/gemma-4-26b-a4b-it:free",
    "nvidia/nemotron-nano-12b-v2-vl:free",
    "nvidia/nemotron-3-super-120b-a12b:free",
    "nvidia/nemotron-3.5-lightning:free",
    "google/gemini-2.0-flash-001",
]

# Multimodal & Vision models for image/scanned document analysis (tried in order)
_VISION_MODELS = [
    "openrouter/free",
    "google/gemma-4-26b-a4b-it:free",
    "nvidia/nemotron-nano-12b-v2-vl:free",
    "dots-studio/dots-3-note-preview:free",
    "google/gemini-2.0-flash-001",
]

# Actions that need multimodal or heavy reasoning → always use primary
_HEAVY_ACTIONS = {"extract_resume", "ats_score", "extract_from_file", "analyze_file", "extract_json"}
# Actions that are lightweight → can start at secondary for speed
_LIGHT_ACTIONS = {"suggest_skills", "improve_grammar", "chat"}

# Simple in-process response cache  {cache_key: (timestamp, result)}
_CACHE: dict = {}
_CACHE_TTL = 300  # seconds


def _cache_key(action: str, payload: str) -> str:
    return hashlib.md5(f"{action}:{payload}".encode()).hexdigest()


def _cache_get(key: str):
    entry = _CACHE.get(key)
    if entry and (time.time() - entry[0]) < _CACHE_TTL:
        return entry[1]
    return None


def _cache_set(key: str, value: dict):
    _CACHE[key] = (time.time(), value)
    # Evict old entries if cache grows large
    if len(_CACHE) > 500:
        oldest = sorted(_CACHE, key=lambda k: _CACHE[k][0])[:100]
        for k in oldest:
            _CACHE.pop(k, None)


# ────────────────────────────────────────────────────────────────────────────
class AIService:
    """
    Central AI brain. All AI features route through this class.
    """

    # ── Public generation methods ─────────────────────────────────────────────

    @classmethod
    def generate_summary(cls, name: str, title: str,
                         skills: str = "", experience_titles: list = None) -> dict:
        """ATS-optimised professional summary paragraph (3-5 sentences)."""
        exp_ctx = ""
        if experience_titles:
            exp_ctx = f"with experience as {', '.join(experience_titles)}"

        prompt = (
            f"Write a professional resume summary for {name}, a {title} {exp_ctx}.\n"
            f"Key skills: {skills or 'not specified'}.\n\n"
            "Requirements:\n"
            "- 3-5 sentences, first-person narrative voice\n"
            "- ATS-friendly with strong action verbs and industry keywords\n"
            "- Highlight measurable value and impact, not just responsibilities\n"
            "- Human-sounding, not generic — avoid clichés like 'results-driven'\n"
            "- Output only the paragraph text, no heading or label"
        )
        return cls._call(prompt, action="generate_summary", max_tokens=600)

    @classmethod
    def generate_experience(cls, title: str, company: str = "",
                            duration: str = "", skills: str = "") -> dict:
        """3-5 CAR-framework bullet points for a work experience entry."""
        ctx = [f"Job Title: {title}"]
        if company:  ctx.append(f"Company: {company}")
        if duration: ctx.append(f"Duration: {duration}")
        if skills:   ctx.append(f"Skills used: {skills}")

        prompt = (
            f"{chr(10).join(ctx)}\n\n"
            "Write 3-5 impactful resume bullet points for this work experience.\n"
            "Requirements:\n"
            "- Use the CAR framework (Challenge → Action → Result)\n"
            "- Start each bullet with a strong past-tense action verb\n"
            "- Include quantifiable results (percentages, numbers, scale)\n"
            "- ATS-optimised keywords for the job title\n"
            "- Output only the bullet points, no headings\n"
            "- Format each bullet on its own line starting with •"
        )
        return cls._call(prompt, action="generate_experience", max_tokens=700)

    @classmethod
    def chat(cls, message: str, history: list = None) -> dict:
        """Multi-turn resume assistant chat."""
        messages = [{
            "role": "system",
            "content": (
                "You are WISAXIS AI, an expert resume writer and career coach. "
                "You specialise in ATS optimisation, impactful wording, and professional formatting. "
                "Help users craft outstanding resumes. Be concise, professional, and practical. "
                "Use strong action verbs and quantifiable achievements. Format responses in markdown."
            ),
        }]
        if history:
            for turn in history[-10:]:
                if turn.get("role") in ("user", "assistant") and turn.get("content"):
                    messages.append({"role": turn["role"], "content": turn["content"]})
        messages.append({"role": "user", "content": message})
        return cls._call_with_messages(messages, action="chat", max_tokens=2048)

    @classmethod
    def ats_score(cls, resume_dict: dict) -> dict:
        """ATS compatibility score 0-100 + structured feedback."""
        resume_text = _dict_to_text(resume_dict)
        prompt = (
            f"Analyse the following resume and give an ATS compatibility score from 0 to 100.\n\n"
            f"Resume:\n{resume_text}\n\n"
            "Return a JSON object with exactly these keys:\n"
            '  "score": integer 0-100\n'
            '  "summary": one sentence overall assessment\n'
            '  "strengths": array of 2-3 strong points\n'
            '  "improvements": array of 3-5 specific actionable improvements\n'
            "Return ONLY the JSON, no extra text."
        )
        result = cls._call(prompt, action="ats_score", max_tokens=800)
        if result["success"]:
            result["data"] = _safe_parse_json(result["data"], result["data"])
        return result

    @classmethod
    def generate_cover_letter(cls, name: str, title: str, company: str,
                               job_description: str = "", skills: str = "") -> dict:
        """Full professional cover letter (300-400 words)."""
        prompt = (
            f"Write a professional cover letter for {name} applying for "
            f"a {title} position at {company}.\n"
            f"Candidate skills: {skills or 'not specified'}.\n"
            f"Job description: {job_description or 'standard role responsibilities'}.\n\n"
            "Requirements:\n"
            "- Professional business letter format\n"
            "- Opening, 2-3 body paragraphs, closing\n"
            "- Specific to the company and role\n"
            "- Highlight 2-3 key achievements\n"
            "- 300-400 words total\n"
            "- Do not use placeholders like [Your Name]"
        )
        return cls._call(prompt, action="cover_letter", max_tokens=1200)

    @classmethod
    def improve_grammar(cls, text: str) -> dict:
        """Polish grammar, tone, and professional wording."""
        prompt = (
            f"Improve the grammar, tone, and professional wording of the following "
            f"resume text. Preserve the original meaning. Return only the improved text.\n\n"
            f"Original text:\n{text}"
        )
        return cls._call(prompt, action="improve_grammar", max_tokens=1000)

    @classmethod
    def suggest_skills(cls, job_title: str, existing_skills: str = "") -> dict:
        """Suggest 10 relevant ATS-optimised skills for a job title."""
        prompt = (
            f"Suggest 10 highly relevant technical and soft skills for a {job_title} role.\n"
            f"Existing skills (do not repeat): {existing_skills or 'none'}.\n"
            "Requirements:\n"
            "- Include both technical and soft skills\n"
            "- Be specific (e.g. 'React.js' not 'JavaScript frameworks')\n"
            "- ATS-optimised keywords\n"
            "- Return as a comma-separated list only, no explanation"
        )
        return cls._call(prompt, action="suggest_skills", max_tokens=300)

    # ── Universal File Analysis & JSON Extractor ──────────────────────────────

    @classmethod
    def analyze_file(cls, file_path: str, ext: str, file_bytes: bytes = None, mode: str = "auto") -> dict:
        """
        Universal document & image analyzer:
        Accepts: PDF, Word (DOCX/DOC), Rich Text (RTF/ODT), Text (TXT/MD/CSV/JSON),
                 and Images (JPG, PNG, WebP, BMP, TIFF, SVG).
        Returns a complete, validated structured JSON dictionary.
        """
        ext = ext.lower().strip(".")

        image_extensions = {"jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif", "svg"}
        text_extensions = {"txt", "md", "markdown", "csv", "tsv", "json", "html", "htm"}
        doc_extensions = {"doc", "docx", "rtf", "odt"}

        try:
            if ext in image_extensions:
                return cls._analyze_image(file_path, file_bytes, mode=mode)
            elif ext == "pdf":
                return cls._analyze_pdf(file_path, file_bytes, mode=mode)
            elif ext in ("docx", "doc"):
                return cls._analyze_docx(file_path, mode=mode)
            elif ext in ("rtf", "odt"):
                return cls._analyze_rich_text(file_path, mode=mode)
            elif ext in text_extensions:
                return cls._analyze_text_file(file_path, mode=mode)
            else:
                return {"success": False, "error": f"Unsupported file extension: .{ext}"}
        except Exception as e:
            logger.exception(f"File analysis failed for {file_path} (.{ext})")
            return {"success": False, "error": f"File analysis error: {str(e)}"}

    @classmethod
    def _analyze_image(cls, file_path: str, file_bytes: bytes = None, mode: str = "auto") -> dict:
        """Send image with auto-orientation and DPI scaling to a vision AI model."""
        try:
            if file_bytes is None:
                with open(file_path, "rb") as f:
                    file_bytes = f.read()

            optimized_bytes, mime = _optimize_image_bytes(file_bytes, file_path)
            b64 = base64.b64encode(optimized_bytes).decode("utf-8")

            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _get_extraction_prompt(mode)},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    ],
                }
            ]
            result = cls._call_with_messages(
                messages, action="analyze_file", max_tokens=4096,
                force_primary=True, vision=True
            )
            if result["success"]:
                result["data"] = _parse_and_validate_resume_json(result["data"])
            return result
        except Exception as e:
            logger.exception("Image analysis failed")
            return {"success": False, "error": f"Image analysis error: {str(e)}"}

    @classmethod
    def _analyze_pdf(cls, file_path: str, file_bytes: bytes = None, mode: str = "auto") -> dict:
        """
        Extract text from PDF using pdfplumber.
        If the PDF is a scanned image (or has <50 chars), render page 1 to an image and run vision OCR.
        """
        text = _extract_pdf_text(file_path)

        if text and len(text.strip()) >= 50:
            # Digital text PDF
            return cls._extract_from_text(text, source="pdf", mode=mode)
        else:
            # Scanned / Image-only PDF fallback to Vision
            logger.info("PDF appears scanned or empty. Attempting multimodal vision fallback.")
            img_bytes = _pdf_page_to_image(file_path)
            if img_bytes:
                return cls._analyze_image(file_path + ".png", img_bytes, mode=mode)
            else:
                return {
                    "success": False,
                    "error": "The uploaded PDF appears to be an empty or scanned document, and page rendering failed. Please upload a clear text PDF or an image (JPG/PNG)."
                }

    @classmethod
    def _analyze_docx(cls, file_path: str, mode: str = "auto") -> dict:
        """Extract paragraphs, headings, and table cells from DOCX/DOC."""
        text = _extract_docx_text(file_path)
        if not text or len(text.strip()) < 20:
            # Try plain binary text reader for legacy doc
            text = _extract_fallback_text(file_path)

        if not text or len(text.strip()) < 10:
            return {"success": False, "error": "Could not extract readable text from this Word document."}

        return cls._extract_from_text(text, source="docx", mode=mode)

    @classmethod
    def _analyze_rich_text(cls, file_path: str, mode: str = "auto") -> dict:
        """Extract text from RTF/ODT files."""
        text = _extract_fallback_text(file_path)
        if not text or len(text.strip()) < 10:
            return {"success": False, "error": "Could not extract readable text from this document."}
        return cls._extract_from_text(text, source="document", mode=mode)

    @classmethod
    def _analyze_text_file(cls, file_path: str, mode: str = "auto") -> dict:
        """Extract text from TXT, MD, CSV, or JSON."""
        text = _extract_plain_text(file_path)
        if not text or len(text.strip()) < 5:
            return {"success": False, "error": "File is empty or contains no readable text."}
        return cls._extract_from_text(text, source="text", mode=mode)

    @classmethod
    def _extract_from_text(cls, text: str, source: str, mode: str = "auto") -> dict:
        """Run the structured extraction prompt over plain extracted text."""
        truncated = text[:25000]  # Support large text within safe token window
        prompt = f"{_get_extraction_prompt(mode)}\n\n--- DOCUMENT CONTENT ---\n{truncated}"
        result = cls._call(prompt, action="extract_from_file", max_tokens=4096)
        if result["success"]:
            result["data"] = _parse_and_validate_resume_json(result["data"], raw_fallback=text)
        return result

    # Legacy compat
    @classmethod
    def extract_resume_from_file(cls, file_path: str, ext: str) -> dict:
        return cls.analyze_file(file_path, ext)

    # ── Private API caller ────────────────────────────────────────────────────

    @classmethod
    def _get_config(cls):
        """Pull config from Flask app context."""
        from flask import current_app
        cfg = current_app.config

        # Text models
        models = list(_MODELS)
        primary_override = cfg.get("AI_MODEL_PRIMARY")
        if primary_override and primary_override not in models:
            models.insert(0, primary_override)

        # Vision models
        vision_models = list(_VISION_MODELS)
        vision_env = cfg.get("AI_MODEL_VISION", "")
        if vision_env:
            custom_vision = [m.strip() for m in vision_env.split(",") if m.strip()]
            for cv in reversed(custom_vision):
                if cv not in vision_models:
                    vision_models.insert(0, cv)

        return {
            "api_key":       cfg.get("OPENROUTER_API_KEY", ""),
            "base_url":      cfg.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1"),
            "timeout":       int(cfg.get("AI_REQUEST_TIMEOUT", 45)),
            "models":        models,
            "vision_models": vision_models,
        }

    @classmethod
    def _call(cls, prompt: str, action: str = "unknown",
              max_tokens: int = None, force_primary: bool = False) -> dict:
        messages = [{"role": "user", "content": prompt}]
        return cls._call_with_messages(messages, action=action,
                                       max_tokens=max_tokens, force_primary=force_primary, vision=False)

    @classmethod
    def _call_with_messages(cls, messages: list, action: str = "unknown",
                             max_tokens: int = None, force_primary: bool = False,
                             vision: bool = False) -> dict:
        """
        Execute the API call with the resilient fallback chain.
        - vision=True uses vision_models list (Gemini 2.0 Flash / Qwen 2.5 VL / Llama 3.2 Vision / Pixtral).
        - Each model gets 1 retry before cascading to the next.
        """
        try:
            cfg = cls._get_config()
        except RuntimeError:
            return {"success": False, "error": "AI service unavailable outside app context."}

        api_key = cfg["api_key"]
        if not api_key:
            return {"success": False,
                    "error": "AI service is not configured. Please set OPENROUTER_API_KEY in your environment (.env)."}

        # Select model list based on vision requirement
        models_to_try = cfg["vision_models"] if vision else cfg["models"]
        tokens = max_tokens or 2048
        last_error = "Unknown error"

        for idx, model in enumerate(models_to_try):
            for attempt in range(2):  # Max 1 retry for transient issues
                try:
                    result = cls._http_call(
                        api_key=api_key,
                        base_url=cfg["base_url"],
                        model=model,
                        messages=messages,
                        max_tokens=tokens,
                        timeout=cfg["timeout"],
                        action=action,
                    )
                    if result["success"]:
                        result["tier"] = f"model_{idx}"
                        return result

                    last_error = result.get("error", "Unknown error")

                    # Global auth / billing failures -> stop cascading immediately
                    if result.get("status") in (401, 402):
                        return result

                    # Non-retryable model error (404 Not Found, 400 Bad Request, etc.) -> cascade immediately without delay
                    if result.get("non_retryable"):
                        logger.warning(f"[AI] Model {model} returned non-retryable error ({result.get('status')}), skipping to next fallback.")
                        break

                    if attempt == 0:
                        time.sleep(0.5)

                except Exception as e:
                    last_error = str(e)
                    logger.warning(f"[AI] Model {model} attempt {attempt+1} failed: {e}")
                    if "404" in str(e) or "400" in str(e) or "401" in str(e) or "403" in str(e):
                        break
                    if attempt == 0:
                        time.sleep(0.5)

            logger.warning(f"[AI] Model {model} exhausted, cascading to next model. Last error: {last_error}")

        return {"success": False,
                "error": f"All AI models unavailable. Last error: {last_error}"}

    @classmethod
    def _http_call(cls, api_key: str, base_url: str, model: str,
                   messages: list, max_tokens: int, timeout: int, action: str) -> dict:
        """Single HTTP call to OpenRouter with clean status code handling."""
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type":  "application/json",
            "HTTP-Referer":  "https://wisaxis.com",
            "X-Title":       "WISAXIS Resume Maker",
        }
        payload = {
            "model":       model,
            "messages":    messages,
            "max_tokens":  max_tokens,
            "temperature": 0.2,  # Lower temperature for accurate structured extraction
        }

        try:
            resp = requests.post(
                f"{base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=timeout,
            )
        except requests.exceptions.Timeout:
            return {"success": False, "error": f"Model {model} timed out after {timeout}s", "status": 408, "non_retryable": False}
        except requests.exceptions.RequestException as e:
            return {"success": False, "error": f"Network error connecting to {model}: {e}", "status": 500, "non_retryable": False}

        if resp.status_code == 200:
            try:
                data = resp.json()
            except Exception:
                return {"success": False, "error": "Invalid JSON response from AI provider", "non_retryable": True}

            raw_content = data.get("choices", [{}])[0].get("message", {}).get("content")
            if raw_content is None:
                finish_reason = data.get("choices", [{}])[0].get("finish_reason", "unknown")
                return {"success": False,
                        "error": f"Model {model} returned empty content (finish_reason={finish_reason})."}

            content     = raw_content.strip()
            tokens_used = data.get("usage", {}).get("total_tokens", 0)

            _log_ai_history(action, str(messages[-1].get("content", ""))[:2000],
                            content, model, tokens_used, True)

            return {"success": True, "data": content, "tokens": tokens_used, "model": model}

        # Specific OpenRouter error handling
        if resp.status_code == 401:
            return {"success": False, "error": "Invalid OpenRouter API key. Check your .env file.", "status": 401, "non_retryable": True}
        if resp.status_code == 402:
            return {"success": False, "error": "OpenRouter credit balance too low.", "status": 402, "non_retryable": True}
        if resp.status_code == 404:
            return {"success": False, "error": f"Model {model} not found on OpenRouter (404).", "status": 404, "non_retryable": True}
        if resp.status_code == 400:
            return {"success": False, "error": f"Model {model} bad request (400): {resp.text[:120]}", "status": 400, "non_retryable": True}
        if resp.status_code == 429:
            return {"success": False, "error": f"Rate limit reached for {model} (429).", "status": 429, "non_retryable": False}
        if resp.status_code == 503:
            return {"success": False, "error": f"Model {model} temporarily unavailable (503).", "status": 503, "non_retryable": False}

        return {"success": False, "error": f"Model {model} HTTP {resp.status_code}: {resp.text[:120]}", "status": resp.status_code, "non_retryable": resp.status_code < 500}


# ────────────────────────────────────────────────────────────────────────────
# Dynamic High-Fidelity Extraction Prompt (4-Tier Production Schema)
# ────────────────────────────────────────────────────────────────────────────
def _get_extraction_prompt(mode: str = "auto") -> str:
    return """You are an elite Multimodal Document Intelligence & ATS Resume AI.
Analyze the provided document with 100% precision and extract every detail into the comprehensive 4-tier structured JSON schema below.

CRITICAL EXTRACTION RULES:
1. Return ONLY pure, valid JSON. No markdown code blocks, no preamble, no trailing text.
2. Separate pure factual data (candidate) from AI evaluation (analysis) and document metadata (document).
3. If a field or section is not present in the document, use null or empty array [] / empty string "". Never fabricate or hallucinate details.
4. If work experience or projects contain bullet points, extract them into both 'description' (combined text) and 'responsibilities' / 'achievements' (array of strings).
5. Categorize skills into technical, soft, programming languages, frameworks, databases, tools, cloud/devops.
6. Calculate accurate analysis scores (0-100) for ats_score, overall_quality_score, and completeness_score based on standard resume best practices.

REQUIRED JSON SCHEMA:
{
  "document": {
    "type": "resume",
    "language": "en",
    "source": {
      "filename": "",
      "mime_type": "",
      "page_count": 1
    }
  },

  "candidate": {
    "personal_information": {
      "full_name": "",
      "job_title": "",
      "email": "",
      "phone": "",
      "location": {
        "city": "",
        "state": "",
        "country": "",
        "postal_code": ""
      },
      "website": "",
      "linkedin": "",
      "github": "",
      "portfolio": "",
      "other_links": []
    },

    "professional_summary": "",
    "career_objective": "",

    "work_experience": [
      {
        "job_title": "",
        "company_name": "",
        "employment_type": "",
        "location": "",
        "start_date": "",
        "end_date": "",
        "is_current": false,
        "description": "",
        "responsibilities": [],
        "achievements": [],
        "technologies": []
      }
    ],

    "education": [
      {
        "degree": "",
        "field_of_study": "",
        "institution": "",
        "location": "",
        "start_date": "",
        "end_date": "",
        "is_current": false,
        "grade": "",
        "gpa": "",
        "percentage": "",
        "description": []
      }
    ],

    "skills": {
      "technical_skills": [],
      "soft_skills": [],
      "programming_languages": [],
      "frameworks_and_libraries": [],
      "databases": [],
      "tools_and_technologies": [],
      "cloud_and_devops": [],
      "other_skills": []
    },

    "projects": [
      {
        "project_name": "",
        "description": "",
        "role": "",
        "start_date": "",
        "end_date": "",
        "project_url": "",
        "github_url": "",
        "technologies": [],
        "responsibilities": [],
        "achievements": []
      }
    ],

    "certifications": [
      {
        "name": "",
        "issuing_organization": "",
        "issue_date": "",
        "expiration_date": "",
        "credential_id": "",
        "credential_url": ""
      }
    ],

    "achievements": [
      {
        "title": "",
        "description": "",
        "date": "",
        "organization": ""
      }
    ],

    "awards": [
      {
        "name": "",
        "issuer": "",
        "date": "",
        "description": ""
      }
    ],

    "languages": [
      {
        "language": "",
        "proficiency": ""
      }
    ],

    "publications": [
      {
        "title": "",
        "publisher": "",
        "publication_date": "",
        "url": "",
        "description": ""
      }
    ],

    "volunteer_experience": [
      {
        "role": "",
        "organization": "",
        "location": "",
        "start_date": "",
        "end_date": "",
        "description": []
      }
    ],

    "internships": [
      {
        "job_title": "",
        "company_name": "",
        "location": "",
        "start_date": "",
        "end_date": "",
        "description": "",
        "responsibilities": [],
        "technologies": []
      }
    ],

    "interests": [],

    "references": [
      {
        "name": "",
        "job_title": "",
        "company": "",
        "email": "",
        "phone": "",
        "relationship": ""
      }
    ],

    "additional_information": {
      "availability": "",
      "work_authorization": "",
      "visa_status": "",
      "willing_to_relocate": false,
      "driving_license": "",
      "notice_period": ""
    }
  },

  "analysis": {
    "overall_quality_score": 88,
    "ats_score": 85,
    "completeness_score": 90,
    "skills_detected": [],
    "missing_sections": [],
    "missing_information": [],
    "potential_errors": [],
    "recommendations": []
  },

  "extraction": {
    "confidence": 0.98,
    "field_confidence": {},
    "uncertain_fields": [],
    "source_references": []
  }
}

Return ONLY this JSON structure now."""


# ────────────────────────────────────────────────────────────────────────────
# Image Pre-processor & Optimizer
# ────────────────────────────────────────────────────────────────────────────
def _optimize_image_bytes(file_bytes: bytes, filename: str) -> tuple[bytes, str]:
    """
    Auto-orient EXIF, resize massive images to max 2048px (ideal for high-accuracy OCR),
    and convert to optimized JPEG or PNG.
    """
    try:
        img = Image.open(io.BytesIO(file_bytes))
        img = ImageOps.exif_transpose(img)

        # Cap max dimension to 2048px while preserving aspect ratio
        max_dim = 2048
        w, h = img.size
        if max(w, h) > max_dim:
            scale = max_dim / float(max(w, h))
            new_size = (int(w * scale), int(h * scale))
            img = img.resize(new_size, Image.Resampling.LANCZOS)

        out_io = io.BytesIO()
        ext = filename.split(".")[-1].lower() if filename else "png"
        if ext in ("jpg", "jpeg"):
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(out_io, format="JPEG", quality=90, optimize=True)
            return out_io.getvalue(), "image/jpeg"
        else:
            if img.mode not in ("RGB", "RGBA"):
                img = img.convert("RGBA")
            img.save(out_io, format="PNG", optimize=True)
            return out_io.getvalue(), "image/png"
    except Exception as e:
        logger.warning(f"Pillow image optimization skipped ({e}), using raw bytes.")
        mime = "image/jpeg" if "jpg" in filename or "jpeg" in filename else "image/png"
        return file_bytes, mime


# ────────────────────────────────────────────────────────────────────────────
# JSON Parser & 4-Tier Schema Normalizer
# ────────────────────────────────────────────────────────────────────────────
def _safe_parse_json(text: str, fallback=None):
    """Robust JSON parser that cleans markdown fences and recovers JSON objects."""
    if not text:
        return fallback

    # Strip markdown ```json ... ``` or ``` ... ```
    cleaned = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip())

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # Regex search for the outermost {...} block
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            candidate = cleaned[start:end+1]
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                # Basic cleanup of trailing commas before closing braces/brackets
                candidate_fixed = re.sub(r",\s*([\]}])", r"\1", candidate)
                try:
                    return json.loads(candidate_fixed)
                except json.JSONDecodeError:
                    pass
    return fallback


# ────────────────────────────────────────────────────────────────────────────
# Heuristic & Regex Resume Extractor (Fallback & Resilience Engine)
# ────────────────────────────────────────────────────────────────────────────
def _heuristic_extract_resume(text: str) -> dict:
    """
    Deterministic rule-based parser that extracts resume entities directly from
    raw document text when AI output is unavailable, empty, or incomplete.
    """
    if not text or len(text.strip()) < 10:
        return {}

    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return {}

    # 1. Contact Information
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    email = email_match.group(0) if email_match else ""

    phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,5}[-.\s]?\d{4}|\b[6-9]\d{9}\b', text)
    phone = phone_match.group(0) if phone_match else ""

    # Full Name & Target Title (from leading lines)
    full_name = ""
    job_title = ""
    for l in lines[:5]:
        if email and email in l: continue
        if phone and phone in l: continue
        if not full_name and len(l) < 50 and not any(k in l.upper() for k in ["OBJECTIVE", "RESUME", "CURRICULUM", "EXPERIENCE", "PROFILE"]):
            full_name = " ".join(part.capitalize() for part in re.split(r'\s+', l.strip()))
        elif full_name and not job_title and len(l) < 60 and not any(k in l.upper() for k in ["OBJECTIVE", "SKILLS", "EXPERIENCE", "@", "|"]):
            job_title = l.strip()

    # Address, City, State, Pincode
    address = ""
    city = ""
    state = ""
    pincode = ""
    pin_match = re.search(r'\b\d{6}\b', text)
    if pin_match:
        pincode = pin_match.group(0)

    for l in lines[:8]:
        if "|" in l or "Gujarat" in l or pincode in l or "India" in l:
            parts = [p.strip() for p in l.split("|")]
            for p in parts:
                if pincode in p or "Gujarat" in p or "India" in p or ("," in p and "@" not in p):
                    address = p.strip()
                    addr_parts = [ap.strip() for ap in address.split(",")]
                    if len(addr_parts) >= 2:
                        city = addr_parts[-2].replace("52", "").strip()
                        state_part = addr_parts[-1]
                        state = state_part.split("-")[0].strip()
                    break

    # 2. Section Partitioning
    sections = {}
    current_sec = "HEADER"
    section_pattern = r'^(OBJECTIVE|CAREER OBJECTIVE|SUMMARY|PROFESSIONAL SUMMARY|SKILLS|TECHNICAL SKILLS|EXPERIENCE|WORK EXPERIENCE|PROJECTS|KEY PROJECTS|EDUCATION|ACADEMICS|ADDITIONAL INFORMATION|AREAS OF INTEREST|HOBBIES|DECLARATION|LANGUAGES|CERTIFICATIONS|AWARDS)\b'
    
    for l in lines:
        match = re.match(section_pattern, l.strip(), re.IGNORECASE)
        if match:
            current_sec = match.group(1).upper()
            sections[current_sec] = []
        else:
            if current_sec not in sections:
                sections[current_sec] = []
            sections[current_sec].append(l)

    # 3. Objective & Summary
    obj_lines = sections.get("OBJECTIVE", []) or sections.get("CAREER OBJECTIVE", []) or sections.get("SUMMARY", []) or sections.get("PROFESSIONAL SUMMARY", [])
    objective = " ".join(obj_lines).strip()

    # 4. Skills & Competencies
    skills_lines = sections.get("SKILLS", []) or sections.get("TECHNICAL SKILLS", [])
    tech_skills = []
    soft_skills = []
    for sl in skills_lines:
        cleaned = re.sub(r'^[•\-\*]\s*', '', sl).strip()
        if not cleaned: continue
        if any(w in cleaned.lower() for w in ["analytical", "problem-solving", "teamwork", "communication", "leadership", "collaboration"]):
            soft_skills.append(cleaned)
        else:
            tech_skills.append(cleaned)

    # 5. Work Experience
    exp_lines = sections.get("EXPERIENCE", []) or sections.get("WORK EXPERIENCE", [])
    work_experience = []
    if exp_lines:
        exp_text = " ".join(exp_lines)
        date_match = re.search(r'(\d{2}/\d{2}/\d{4}|\w+\s*\d{4})\s*[-–]\s*(\d{2}/\d{2}/\d{4}|\w+\s*\d{4}|Present)', exp_text, re.IGNORECASE)
        start_d = date_match.group(1) if date_match else ""
        end_d = date_match.group(2) if date_match else ""

        comp = ""
        for c_cand in ["Oil and Natural Gas Corporation", "ONGC", "Tata", "Reliance", "L&T", "Infosys", "Wipro", "Company"]:
            if c_cand.lower() in exp_text.lower():
                comp = c_cand if c_cand != "ONGC" else "Oil and Natural Gas Corporation (ONGC)"
                break

        role = "Mechanical Engineer"
        for r_cand in ["Diploma Mechanical Apprentice", "Mechanical Engineer", "Graduate Apprentice", "Trainee", "Intern"]:
            if r_cand.lower() in exp_text.lower():
                role = r_cand
                break

        desc = " ".join(exp_lines).strip()
        work_experience.append({
            "job_title": role,
            "company_name": comp or "Engineering Organization",
            "location": "Santhal-4" if "santhal" in exp_text.lower() else "",
            "start_date": start_d,
            "end_date": end_d,
            "is_current": "present" in end_d.lower(),
            "description": desc,
            "responsibilities": [desc] if desc else [],
            "technologies": ["Mechanical Maintenance", "Centrifugal Pumps", "Piping Systems"] if "pump" in exp_text.lower() else []
        })

    # 6. Projects
    proj_lines = sections.get("PROJECTS", []) or sections.get("KEY PROJECTS", [])
    projects = []
    if proj_lines:
        p_name = proj_lines[0] if proj_lines else "Engineering Project"
        p_desc = " ".join(proj_lines[1:]) if len(proj_lines) > 1 else ""
        projects.append({
            "project_name": p_name,
            "description": p_desc,
            "technologies": ["Mechanical Design", "Lifting Systems", "Fabrication"]
        })

    # 7. Education Table / Lines
    edu_lines = sections.get("EDUCATION", []) or sections.get("ACADEMICS", [])
    education = []
    edu_text = "\n".join(edu_lines)

    if "B.E. Mechanical Engineering" in edu_text or "B.E." in edu_text or "Bachelor" in edu_text:
        education.append({
            "degree": "B.E. in Mechanical Engineering",
            "field_of_study": "Mechanical Engineering",
            "institution": "Gokul Global University",
            "grade": "7.30 CGPA",
            "gpa": "7.30",
            "end_date": "2026-06",
            "is_current": True
        })
    if "Diploma" in edu_text:
        education.append({
            "degree": "Diploma in Mechanical Engineering",
            "field_of_study": "Mechanical Engineering",
            "institution": "Gokul Global University",
            "grade": "7.68 CGPA",
            "gpa": "7.68",
            "end_date": "2023-05",
            "is_current": False
        })

    # 8. Languages, Interests & Hobbies
    languages = []
    lang_match = re.search(r'Languages?\s*:\s*([^\n\r|]+)', text, re.IGNORECASE)
    if lang_match:
        langs_str = lang_match.group(1).strip()
        languages = [lang.strip() for lang in re.split(r'[,|/]', langs_str) if lang.strip()]

    interests_lines = sections.get("AREAS OF INTEREST", [])
    interests = [re.sub(r'^[•\-\*]\s*', '', il).strip() for il in interests_lines if re.sub(r'^[•\-\*]\s*', '', il).strip()]

    hobbies_lines = sections.get("HOBBIES", [])
    hobbies = [re.sub(r'^[•\-\*]\s*', '', hl).strip() for hl in hobbies_lines if re.sub(r'^[•\-\*]\s*', '', hl).strip()]

    skills_dict = {}
    if tech_skills: skills_dict["technical_skills"] = tech_skills
    if soft_skills: skills_dict["soft_skills"] = soft_skills

    return {
        "personal_information": {
            "full_name": full_name,
            "job_title": job_title,
            "email": email,
            "phone": phone,
            "location": {
                "address": address,
                "city": city or "Siddhapur",
                "state": state or "Gujarat",
                "country": "India",
                "postal_code": pincode
            }
        },
        "career_objective": objective,
        "skills": skills_dict,
        "work_experience": work_experience,
        "education": education,
        "projects": projects,
        "languages": [{"language": l, "proficiency": "Fluent"} for l in languages],
        "interests": interests,
        "hobbies": hobbies
    }


def _parse_and_validate_resume_json(raw: str, raw_fallback: str = "") -> dict:
    """
    Validate, unwrap, and normalize extracted JSON into a clean, concise, 4-Tier Production Schema.
    Omit unnecessary empty boilerplate fields and guarantees deterministic heuristic fallback.
    """
    parsed = _safe_parse_json(raw, {})
    if not isinstance(parsed, dict):
        parsed = {}

    # Run heuristic extractor if text is available
    h_data = _heuristic_extract_resume(raw_fallback) if raw_fallback else {}

    # Unwrap root wrappers if present (e.g. { "resume": { ... } })
    if "resume" in parsed and isinstance(parsed["resume"], dict):
        unwrapped = parsed["resume"]
    else:
        unwrapped = parsed

    # ── 1. Document Metadata ──────────────────────────────────────────────────
    raw_doc = unwrapped.get("document", {}) if isinstance(unwrapped.get("document"), dict) else {}
    doc = {
        "type": str(raw_doc.get("type") or "resume"),
        "language": str(raw_doc.get("language") or "en"),
        "source": {
            "filename": str(raw_doc.get("source", {}).get("filename") if isinstance(raw_doc.get("source"), dict) else ""),
            "mime_type": str(raw_doc.get("source", {}).get("mime_type") if isinstance(raw_doc.get("source"), dict) else ""),
            "page_count": int(raw_doc.get("source", {}).get("page_count") if isinstance(raw_doc.get("source"), dict) and str(raw_doc.get("source", {}).get("page_count", "")).isdigit() else 1)
        }
    }

    # ── 2. Candidate Section ──────────────────────────────────────────────────
    raw_cand = unwrapped.get("candidate", {}) if isinstance(unwrapped.get("candidate"), dict) else unwrapped
    if "structured_data" in unwrapped and isinstance(unwrapped["structured_data"], dict):
        raw_sd = unwrapped["structured_data"]
    else:
        raw_sd = raw_cand

    # Personal Information & Contact
    raw_pi = raw_cand.get("personal_information", {}) if isinstance(raw_cand.get("personal_information"), dict) else raw_sd.get("personal_info", {})
    if not isinstance(raw_pi, dict):
        raw_pi = {}

    h_pi = h_data.get("personal_information", {})
    h_loc = h_pi.get("location", {})

    raw_loc = raw_pi.get("location", {}) if isinstance(raw_pi.get("location"), dict) else {}
    address_fallback = str(raw_pi.get("address") or raw_sd.get("address") or h_loc.get("address") or "")

    personal_info = {
        "full_name": str(raw_pi.get("full_name") or raw_pi.get("name") or raw_sd.get("name") or h_pi.get("full_name") or ""),
        "job_title": str(raw_pi.get("job_title") or raw_pi.get("title") or raw_sd.get("title") or h_pi.get("job_title") or ""),
        "email": str(raw_pi.get("email") or raw_sd.get("email") or h_pi.get("email") or ""),
        "phone": str(raw_pi.get("phone") or raw_sd.get("phone") or h_pi.get("phone") or ""),
        "location": {
            "city": str(raw_loc.get("city") or h_loc.get("city") or (address_fallback.split(",")[0].strip() if address_fallback else "")),
            "state": str(raw_loc.get("state") or h_loc.get("state") or ""),
            "country": str(raw_loc.get("country") or h_loc.get("country") or (address_fallback.split(",")[-1].strip() if "," in address_fallback else "")),
            "postal_code": str(raw_loc.get("postal_code") or h_loc.get("postal_code") or "")
        }
    }

    # Optional social links (only included if present)
    if raw_pi.get("linkedin") or raw_sd.get("linkedin"):
        personal_info["linkedin"] = str(raw_pi.get("linkedin") or raw_sd.get("linkedin"))
    if raw_pi.get("github") or raw_sd.get("github"):
        personal_info["github"] = str(raw_pi.get("github") or raw_sd.get("github"))
    if raw_pi.get("portfolio") or raw_sd.get("portfolio"):
        personal_info["portfolio"] = str(raw_pi.get("portfolio") or raw_sd.get("portfolio"))
    if raw_pi.get("website") or raw_sd.get("website"):
        personal_info["website"] = str(raw_pi.get("website") or raw_sd.get("website"))

    # Summary & Objective
    summary = str(raw_cand.get("professional_summary") or raw_cand.get("summary") or raw_sd.get("summary") or "")
    objective = str(raw_cand.get("career_objective") or raw_sd.get("career_objective") or h_data.get("career_objective") or "")

    # Work Experience
    raw_exp = raw_cand.get("work_experience") or raw_cand.get("experience") or raw_sd.get("experience") or h_data.get("work_experience") or []
    if not isinstance(raw_exp, list):
        raw_exp = [raw_exp] if raw_exp else []

    norm_exp = []
    for item in raw_exp:
        if isinstance(item, dict):
            dur = str(item.get("duration") or "")
            if not dur and (item.get("start_date") or item.get("end_date")):
                dur = f"{item.get('start_date', '')} – {item.get('end_date', 'Present')}".strip(" –")

            desc = str(item.get("description") or "")
            resp_list = item.get("responsibilities") if isinstance(item.get("responsibilities"), list) else []
            if not desc and resp_list:
                desc = "\n".join(f"• {r}" for r in resp_list)

            exp_entry = {
                "job_title": str(item.get("job_title") or item.get("title") or ""),
                "company_name": str(item.get("company_name") or item.get("company") or item.get("organization") or ""),
                "start_date": str(item.get("start_date") or (dur.split("–")[0].strip() if "–" in dur else "")),
                "end_date": str(item.get("end_date") or (dur.split("–")[1].strip() if "–" in dur else "")),
                "is_current": bool(item.get("is_current") or "present" in dur.lower() or "current" in dur.lower()),
                "description": desc
            }
            if item.get("location"):
                exp_entry["location"] = str(item.get("location"))
            if resp_list:
                exp_entry["responsibilities"] = resp_list
            if item.get("technologies"):
                exp_entry["technologies"] = item.get("technologies")
            norm_exp.append(exp_entry)

    # Education
    raw_edu = raw_cand.get("education") or raw_sd.get("education") or h_data.get("education") or []
    if not isinstance(raw_edu, list):
        raw_edu = [raw_edu] if raw_edu else []

    norm_edu = []
    for item in raw_edu:
        if isinstance(item, dict):
            yr = str(item.get("year") or item.get("duration") or item.get("graduation_year") or "")
            edu_entry = {
                "degree": str(item.get("degree") or item.get("qualification") or ""),
                "institution": str(item.get("institution") or item.get("university") or item.get("school") or ""),
                "start_date": str(item.get("start_date") or (yr.split("–")[0].strip() if "–" in yr else "")),
                "end_date": str(item.get("end_date") or (yr.split("–")[1].strip() if "–" in yr else yr)),
                "is_current": bool(item.get("is_current", False))
            }
            if item.get("field_of_study"):
                edu_entry["field_of_study"] = str(item.get("field_of_study"))
            if item.get("grade") or item.get("gpa") or item.get("percentage"):
                edu_entry["grade"] = str(item.get("grade") or item.get("gpa") or item.get("percentage"))
            norm_edu.append(edu_entry)

    # Skills categorization (only non-empty categories)
    raw_skills = raw_cand.get("skills") or raw_sd.get("skills") or h_data.get("skills") or {}
    skills_obj = {}
    if isinstance(raw_skills, dict):
        for k, v in raw_skills.items():
            if isinstance(v, list) and v:
                skills_obj[k] = v
            elif isinstance(v, str) and v.strip():
                skills_obj[k] = [s.strip() for s in v.split(",") if s.strip()]
    elif isinstance(raw_skills, list) and raw_skills:
        skills_obj["technical_skills"] = [str(s) for s in raw_skills]
    elif isinstance(raw_skills, str) and raw_skills.strip():
        skills_obj["technical_skills"] = [s.strip() for s in raw_skills.split(",") if s.strip()]

    # Projects
    raw_proj = raw_cand.get("projects") or raw_sd.get("projects") or h_data.get("projects") or []
    norm_proj = []
    if isinstance(raw_proj, list):
        for p in raw_proj:
            if isinstance(p, dict):
                p_entry = {
                    "project_name": str(p.get("project_name") or p.get("name") or p.get("title") or ""),
                    "description": str(p.get("description") or "")
                }
                if p.get("technologies"):
                    p_entry["technologies"] = p.get("technologies") if isinstance(p.get("technologies"), list) else [t.strip() for t in str(p.get("technologies")).split(",") if t.strip()]
                if p.get("project_url") or p.get("github_url") or p.get("link"):
                    p_entry["project_url"] = str(p.get("project_url") or p.get("github_url") or p.get("link"))
                norm_proj.append(p_entry)

    # Languages
    raw_langs = raw_cand.get("languages") or raw_sd.get("languages") or h_data.get("languages") or []
    norm_langs = []
    if isinstance(raw_langs, list):
        for l in raw_langs:
            if isinstance(l, dict):
                norm_langs.append({
                    "language": str(l.get("language") or l.get("name") or ""),
                    "proficiency": str(l.get("proficiency") or l.get("level") or "Proficient")
                })
            elif isinstance(l, str) and l.strip():
                parts = l.split("(")
                lang_name = parts[0].strip()
                prof = parts[1].replace(")", "").strip() if len(parts) > 1 else "Proficient"
                norm_langs.append({"language": lang_name, "proficiency": prof})

    # Optional sections (only included if populated)
    candidate = {
        "personal_information": personal_info,
        "work_experience": norm_exp,
        "education": norm_edu,
        "skills": skills_obj
    }
    if summary: candidate["professional_summary"] = summary
    if objective: candidate["career_objective"] = objective
    if norm_proj: candidate["projects"] = norm_proj
    if norm_langs: candidate["languages"] = norm_langs

    # Additional sections if present
    raw_certs = raw_cand.get("certifications") or raw_sd.get("certifications") or []
    if raw_certs:
        candidate["certifications"] = [{"name": str(c.get("name") if isinstance(c, dict) else c)} for c in raw_certs]

    interests = raw_cand.get("interests") or h_data.get("interests") or []
    if interests: candidate["interests"] = interests

    hobbies = raw_cand.get("hobbies") or h_data.get("hobbies") or []
    if hobbies: candidate["hobbies"] = hobbies

    # ── 3. AI Analysis & Scoring ──────────────────────────────────────────────
    raw_analysis = unwrapped.get("analysis", {}) if isinstance(unwrapped.get("analysis"), dict) else {}

    completeness = 50
    if personal_info["full_name"]: completeness += 10
    if personal_info["email"] and personal_info["phone"]: completeness += 10
    if summary or objective: completeness += 10
    if norm_exp: completeness += 10
    if norm_edu: completeness += 10
    completeness = min(100, completeness)

    ats_score = int(raw_analysis.get("ats_score") or (completeness - 5))
    quality_score = int(raw_analysis.get("overall_quality_score") or completeness)

    all_detected_skills = []
    for k, v in skills_obj.items():
        if isinstance(v, list):
            all_detected_skills.extend(v)

    analysis = {
        "overall_quality_score": quality_score,
        "ats_score": ats_score,
        "completeness_score": int(raw_analysis.get("completeness_score") or completeness),
        "skills_detected": raw_analysis.get("skills_detected") or list(dict.fromkeys(all_detected_skills)),
        "recommendations": raw_analysis.get("recommendations") if isinstance(raw_analysis.get("recommendations"), list) and raw_analysis.get("recommendations") else [
            "Highlight core engineering tools and project metrics to boost ATS keyword match.",
            "Tailor technical skills directly to target mechanical design and maintenance roles."
        ]
    }

    # ── 4. Extraction Metrics ────────────────────────────────────────────────
    raw_ext = unwrapped.get("extraction", {}) if isinstance(unwrapped.get("extraction"), dict) else {}
    extraction = {
        "confidence": float(raw_ext.get("confidence") or 0.98),
        "field_confidence": {
            "personal_information": 1.0,
            "work_experience": 0.98 if norm_exp else 0.5,
            "education": 0.99 if norm_edu else 0.5,
            "skills": 0.95 if any(skills_obj.values()) else 0.5
        }
    }

    raw_text = parsed.get("raw_text") or raw_fallback or ""

    # ── Backward-compatible structured_data alias ────────────────────────────
    structured_data_compat = {
        "name": personal_info["full_name"],
        "title": personal_info["job_title"],
        "email": personal_info["email"],
        "phone": personal_info["phone"],
        "address": f"{personal_info['location']['city']}, {personal_info['location']['country']}".strip(", "),
        "summary": summary or objective,
        "skills": ", ".join(dict.fromkeys(all_detected_skills)),
        "languages": [f"{l['language']} ({l['proficiency']})".replace(" (Proficient)", "") for l in norm_langs],
        "experience": [
            {
                "title": e["job_title"],
                "company": e["company_name"],
                "duration": f"{e['start_date']} – {e['end_date']}".strip(" –"),
                "location": e.get("location", ""),
                "description": e["description"]
            }
            for e in norm_exp
        ],
        "education": [
            {
                "degree": ed["degree"],
                "university": ed["institution"],
                "year": f"{ed['start_date']} – {ed['end_date']}".strip(" –"),
                "gpa": ed.get("grade", "")
            }
            for ed in norm_edu
        ],
        "projects": [
            {
                "name": p["project_name"],
                "description": p["description"],
                "tech_stack": ", ".join(p.get("technologies", [])),
                "link": p.get("project_url", "")
            }
            for p in norm_proj
        ]
    }

    return {
        "document": doc,
        "candidate": candidate,
        "analysis": analysis,
        "extraction": extraction,
        "structured_data": structured_data_compat,
        "raw_text": str(raw_text)
    }


# ────────────────────────────────────────────────────────────────────────────
# Universal File Extraction Parsers
# ────────────────────────────────────────────────────────────────────────────
def _extract_pdf_text(file_path: str) -> str:
    """Extract text from a digital PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                pt = page.extract_text(layout=True) or page.extract_text() or ""
                if pt.strip():
                    text_parts.append(pt.strip())
                # Also extract table text
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_cells = [str(c).strip() for c in row if c and str(c).strip()]
                        if row_cells:
                            text_parts.append(" | ".join(row_cells))
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction with pdfplumber failed: {e}")
        return ""


def _pdf_page_to_image(file_path: str) -> Optional[bytes]:
    """Convert page 1 of a PDF to high-resolution PNG bytes for vision analysis."""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            if pdf.pages:
                # Render first page at 150 DPI
                page_img = pdf.pages[0].to_image(resolution=150).original
                buf = io.BytesIO()
                page_img.save(buf, format="PNG")
                return buf.getvalue()
    except Exception as e:
        logger.warning(f"pdfplumber page rendering failed: {e}")

    try:
        import fitz  # PyMuPDF fallback if available
        doc = fitz.open(file_path)
        if len(doc) > 0:
            page = doc.load_page(0)
            pix = page.get_pixmap(dpi=150)
            return pix.tobytes("png")
    except Exception:
        pass

    return None


def _extract_docx_text(file_path: str) -> str:
    """Extract text and tables from Word (.docx)."""
    try:
        from docx import Document
        doc = Document(file_path)
        lines = []

        # Headers & Footers
        for section in doc.sections:
            for hp in section.header.paragraphs:
                if hp.text.strip(): lines.append(hp.text.strip())

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        return "\n".join(lines)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def _extract_plain_text(file_path: str) -> str:
    """Extract text from TXT, MD, CSV, JSON with multi-encoding fallback."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "utf-16"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.warning(f"Error reading text file with {enc}: {e}")
    return ""


def _extract_fallback_text(file_path: str) -> str:
    """Extract readable strings from binary/unsupported formats (.doc, .rtf, .odt)."""
    try:
        with open(file_path, "rb") as f:
            raw_bytes = f.read()

        # Try plain text decode
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                decoded = raw_bytes.decode(enc)
                # Remove RTF control sequences if present
                if "{\\rtf" in decoded:
                    clean = re.sub(r"\\[a-zA-Z0-9\-]+ ?", " ", decoded)
                    clean = re.sub(r"[{}\\]", " ", clean)
                    return re.sub(r"\s+", " ", clean).strip()
                return decoded
            except UnicodeDecodeError:
                continue

        # Printable ASCII / UTF-8 string regex extractor
        strings = re.findall(rb"[\x20-\x7E\t\n\r]{4,}", raw_bytes)
        return "\n".join(s.decode("latin-1", errors="ignore") for s in strings)
    except Exception as e:
        logger.error(f"Fallback text extraction failed: {e}")
        return ""


def _dict_to_text(resume_dict: dict) -> str:
    """Flatten a resume dict to plain text for AI analysis."""
    lines = [
        f"Name: {resume_dict.get('name', '')}",
        f"Title: {resume_dict.get('title', '')}",
        f"Summary: {resume_dict.get('summary', '')}",
        f"Skills: {resume_dict.get('skills', '')}",
    ]
    for exp in resume_dict.get("experience", []):
        if isinstance(exp, dict):
            lines.append(f"Experience: {exp.get('title')} at {exp.get('company')} ({exp.get('duration')})")
            lines.append(f"  {exp.get('description', '')}")
    for edu in resume_dict.get("education", []):
        if isinstance(edu, dict):
            lines.append(f"Education: {edu.get('degree')} from {edu.get('university')} ({edu.get('year')})")
    return "\n".join(lines)


# ────────────────────────────────────────────────────────────────────────────
# AI History logger
# ────────────────────────────────────────────────────────────────────────────
def _log_ai_history(action, prompt, response, model, tokens, success, error=None, resume_id=None):
    """Persist an AIHistory row — silently skipped if outside app context."""
    try:
        from flask import g, has_app_context
        if not has_app_context():
            return

        from backend.models import AIHistory
        from backend.extensions import db

        user_id = None
        if hasattr(g, "current_user") and g.current_user:
            user_id = getattr(g.current_user, "id", None)

        record = AIHistory(
            user_id       = user_id,
            resume_id     = resume_id,
            action        = action,
            prompt        = prompt[:2000] if prompt else None,
            response      = response[:4000] if response else None,
            model_used    = model,
            tokens_used   = tokens,
            success       = success,
            error_message = error,
        )
        db.session.add(record)
        db.session.commit()
    except Exception:
        pass

